"""
Procesa archivos .docx: extrae texto, describe imágenes con Claude Vision,
y genera .txt limpios para indexar en ChromaDB.
"""

import os
import sys
import base64
import zipfile
import tempfile
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
import anthropic

KNOWLEDGE_BASE = Path(__file__).parent.parent / "knowledge_base"
client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))


def describe_image_with_claude(image_bytes: bytes, image_ext: str, context: str = "") -> str:
    """Usa Claude Vision para describir una imagen del libro."""
    media_type_map = {
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "gif": "image/gif",
        "webp": "image/webp",
        "bmp": "image/png",  # convertir bmp como png
        "tiff": "image/png",
        "tif": "image/png",
        "wmf": None,  # skip
        "emf": None,  # skip
    }
    media_type = media_type_map.get(image_ext.lower())
    if not media_type:
        return f"[Imagen no soportada: formato {image_ext}]"

    img_b64 = base64.standard_b64encode(image_bytes).decode("utf-8")
    context_text = f"\nContexto del libro: {context}" if context else ""

    try:
        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=1024,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": img_b64,
                            },
                        },
                        {
                            "type": "text",
                            "text": (
                                f"Esta imagen proviene de un libro de coaching para pilotos de aviación "
                                f"escrito por Jaime Ferrer Vives.{context_text}\n\n"
                                "Describe detalladamente esta imagen en español: qué muestra, qué conceptos "
                                "representa, qué información contiene (texto visible, diagramas, tablas, "
                                "gráficos, etc.). Sé específico y completo para que la descripción sea útil "
                                "como texto de búsqueda en un sistema RAG."
                            ),
                        },
                    ],
                }
            ],
        )
        return response.content[0].text
    except Exception as e:
        return f"[Error describiendo imagen: {e}]"


def extract_images_from_docx(docx_path: Path) -> dict:
    """Extrae imágenes del .docx (que es un ZIP) y retorna {rel_id: (bytes, ext)}."""
    images = {}
    with zipfile.ZipFile(docx_path, "r") as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                ext = name.rsplit(".", 1)[-1].lower()
                images[name] = (z.read(name), ext)
    return images


def build_rel_to_image_map(docx_path: Path, images: dict) -> dict:
    """Construye mapa de relationship ID a imagen: {rId: (bytes, ext)}."""
    rel_map = {}
    with zipfile.ZipFile(docx_path, "r") as z:
        if "word/_rels/document.xml.rels" in z.namelist():
            import xml.etree.ElementTree as ET
            rels_xml = z.read("word/_rels/document.xml.rels")
            root = ET.fromstring(rels_xml)
            for rel in root:
                rid = rel.get("Id")
                target = rel.get("Target", "")
                full_target = "word/" + target if not target.startswith("word/") else target
                if full_target in images:
                    rel_map[rid] = images[full_target]
    return rel_map


def get_paragraph_text_and_images(para, rel_map: dict, image_counter: list) -> list:
    """
    Extrae el contenido de un párrafo: texto e imágenes inline.
    Retorna lista de strings (texto o descripción de imagen).
    """
    parts = []
    full_text = para.text.strip()

    # Buscar imágenes inline en el XML del párrafo
    for drawing in para._element.findall(".//" + qn("w:drawing")):
        blip = drawing.find(".//" + qn("a:blip"))
        if blip is not None:
            rid = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if rid and rid in rel_map:
                img_bytes, ext = rel_map[rid]
                image_counter[0] += 1
                n = image_counter[0]
                print(f"  -> Describiendo imagen #{n} (rId={rid}, formato={ext}, tamaño={len(img_bytes)} bytes)...")
                context = full_text if full_text else ""
                desc = describe_image_with_claude(img_bytes, ext, context)
                parts.append(f"\n[IMAGEN #{n}: {desc}]\n")

    # Añadir el texto del párrafo
    if full_text:
        parts.append(full_text)

    return parts


def process_docx(docx_path: Path) -> str:
    """Procesa un .docx completo y retorna el texto limpio con descripciones de imágenes."""
    print(f"\nProcesando: {docx_path.name}")
    doc = Document(str(docx_path))

    images = extract_images_from_docx(docx_path)
    rel_map = build_rel_to_image_map(docx_path, images)
    print(f"  Imágenes encontradas en el archivo: {len(images)}")
    print(f"  Relaciones mapeadas: {len(rel_map)}")

    image_counter = [0]
    output_parts = []

    for para in doc.paragraphs:
        parts = get_paragraph_text_and_images(para, rel_map, image_counter)
        output_parts.extend(parts)

    # También procesar tablas
    for table in doc.tables:
        output_parts.append("\n[TABLA]")
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                output_parts.append(" | ".join(row_texts))
        output_parts.append("[/TABLA]\n")

    result = "\n".join(output_parts)
    # Limpiar líneas vacías excesivas
    import re
    result = re.sub(r'\n{4,}', '\n\n\n', result)
    return result


def main():
    docx_files = [
        "Nuevo formato A5 ver 1 Enero 2026 Color.docx",
        "Las Competencias del Instructor de aviación_final - copia_capitulo7.docx",
    ]

    if len(sys.argv) > 1:
        docx_files = sys.argv[1:]

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("ERROR: Variable de entorno ANTHROPIC_API_KEY no encontrada.")
        sys.exit(1)

    for filename in docx_files:
        docx_path = KNOWLEDGE_BASE / filename
        if not docx_path.exists():
            print(f"ADVERTENCIA: No se encontró {docx_path}")
            continue

        text = process_docx(docx_path)

        out_name = docx_path.stem + ".txt"
        out_path = KNOWLEDGE_BASE / out_name
        out_path.write_text(text, encoding="utf-8")

        word_count = len(text.split())
        char_count = len(text)
        print(f"\n✓ Guardado: {out_path.name}")
        print(f"  Palabras: {word_count:,} | Caracteres: {char_count:,}")


if __name__ == "__main__":
    main()
